"""보고서 JSON → Word(.docx) 변환."""
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x1F, 0x38, 0x64)
FONT = "맑은 고딕"


def _set_korean_font(style, size=None, bold=None, color=None):
    style.font.name = FONT
    # 한글 폰트는 eastAsia 속성도 함께 지정해야 적용된다
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)
    if size:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color:
        style.font.color.rgb = color


def build_docx(report: dict) -> bytes:
    doc = Document()
    _set_korean_font(doc.styles["Normal"], size=10.5)
    for h in ("Heading 1", "Heading 2", "Title"):
        try:
            _set_korean_font(doc.styles[h], color=NAVY)
        except KeyError:
            pass

    # 표지
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(report.get("title", "조사 보고서"))
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = NAVY
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(report.get("subtitle", ""))
    run.font.size = Pt(13)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    doc.add_paragraph()

    # 경영진 요약
    doc.add_heading("Executive Summary", level=1)
    for para in str(report.get("executive_summary", "")).split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # 핵심 발견사항
    if report.get("key_findings"):
        doc.add_heading("핵심 발견사항", level=1)
        for item in report["key_findings"]:
            doc.add_paragraph(str(item), style="List Number")

    # 본문 섹션
    for sec in report.get("sections", []):
        doc.add_heading(sec.get("heading", ""), level=1)
        for para in str(sec.get("content", "")).split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        for bullet in sec.get("bullets") or []:
            doc.add_paragraph(str(bullet), style="List Bullet")

    # 데이터 표
    for table_data in report.get("data_tables", []):
        doc.add_heading(table_data.get("title", "데이터"), level=1)
        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])
        if not headers:
            continue
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for c, h in enumerate(headers):
            cell = table.rows[0].cells[c]
            cell.text = str(h)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
        for row in rows:
            cells = table.add_row().cells
            for c in range(len(headers)):
                cells[c].text = str(row[c]) if c < len(row) else ""

    # 제언
    if report.get("recommendations"):
        doc.add_heading("제언 (Recommendations)", level=1)
        for item in report["recommendations"]:
            doc.add_paragraph(str(item), style="List Number")

    # 출처
    if report.get("sources"):
        doc.add_heading("출처 (Sources)", level=1)
        for s in report["sources"]:
            doc.add_paragraph(
                f"{s.get('title', '')} — {s.get('url', '')}", style="List Bullet"
            )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
